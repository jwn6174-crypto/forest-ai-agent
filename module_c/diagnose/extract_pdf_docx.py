"""
extract_pdf_docx.py — 사용자 제공 4개 reference 자료의 텍스트 추출.

출력: _workspace/analysis/extracted/*.txt
"""

import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[2]
BASE = ROOT / "기반채팅"
OUT = ROOT / "_workspace" / "analysis" / "extracted"
OUT.mkdir(parents=True, exist_ok=True)

FILES = [
    "2025년 산림소득분야 사업시행지침.pdf",
    "★ 2024년 임산물생산조사 보고서.pdf",
    "OpenAPI활용가이드_산림청_산림자원통계 서비스_v1.2.docx",
    "오픈API 활용자가이드_금융위원회_일반상품시세정보.docx",
]


def extract_pdf(pdf_path: Path) -> str:
    import pdfplumber
    parts = []
    with pdfplumber.open(pdf_path) as pdf:
        for i, page in enumerate(pdf.pages):
            t = page.extract_text() or ""
            parts.append(f"--- Page {i+1} ---\n{t}\n")
    return "\n".join(parts)


def extract_docx(docx_path: Path) -> str:
    from docx import Document
    doc = Document(docx_path)
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    # 표 추출
    for ti, table in enumerate(doc.tables):
        parts.append(f"\n[Table {ti+1}]")
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(parts)


def main():
    for fname in FILES:
        path = BASE / fname
        if not path.exists():
            print(f"❌ NOT FOUND: {path}")
            continue
        print(f"\n📄 {fname} ({path.stat().st_size//1024} KB)")
        try:
            if fname.endswith(".pdf"):
                text = extract_pdf(path)
            elif fname.endswith(".docx"):
                text = extract_docx(path)
            else:
                continue
            out_name = fname.replace(" ", "_").replace("★", "star").replace("/", "_") + ".txt"
            (OUT / out_name).write_text(text, encoding="utf-8")
            print(f"   ✅ saved: {OUT / out_name}")
            print(f"   length: {len(text):,} chars")
            print(f"   preview: {text[:300]}")
        except Exception as e:
            print(f"   ❌ error: {e}")


if __name__ == "__main__":
    main()
